# coding=utf-8
import os
from pymongo import MongoClient
from docx import Document

filepath = 'E:/data/科技馆展品记录汇总表.docx'
client = MongoClient("mongodb://serp:serp123456@127.0.0.1:27017")
db = client.kxjy

document = Document(filepath) #读入文件

tables = document.tables
table_nums = len(document.tables)
table_name = ['上海科技馆', '上海自然博物馆', '中国科技馆', '北京老牛儿童馆', '辽宁科技馆', '广西科技馆', '浙江省科技馆', '杭州低碳馆', '绍兴科技馆', '温州科技馆', '重庆科技馆', '合肥科技馆', '厦门科技馆']


for index, table in enumerate(tables):
    datas = []
    venueName = table_name[index]
    print(venueName, " start reading table")
    for i in range(2, len(table.rows)):  # 从表格第2行开始循环读取表格数据 len(table.rows)
        print(venueName, "'s ", i-1)
        n = 1
        xh = table.cell(i, n - 1).text
        zpmc = table.cell(i, n).text
        jbnr = table.cell(i, n + 1).text
        kbdjbm = table.cell(i, n + 2).text
        kbdjqd = table.cell(i, n + 3).text
        rzsp = table.cell(i, n + 4).text
        kxtjys = table.cell(i, n + 5).text
        stseys = table.cell(i, n + 6).text
        kjbj = table.cell(i, n + 7).text
        fzjjxs = table.cell(i, n + 8).text
        fzjjnr = table.cell(i, n + 9).text
        zplb = table.cell(i, n + 10).text
        zcfl = table.cell(i, n + 11).text
        zpjs = table.cell(i, n + 12).text
        jbczykxyyff = table.cell(i, n + 13).text
        xxxg = table.cell(i, n + 14).text
        cgrs = table.cell(i, n + 15).text
        ztmc = table.cell(i, n + 16).text
        ztxsjg = table.cell(i, n + 17).text
        xq = table.cell(i, n + 18).text
        jlsj = table.cell(i, n + 19).text
        jlr = table.cell(i, n + 20).text

        item = {"name": zpmc, "content": jbnr, "standard_code": kbdjbm, "standard_strength": kbdjqd,
                "congnition_level": rzsp, "explore_factor": kxtjys, "STSE_factor": stseys, "space_layout": kjbj,
                "teach_style": fzjjxs, "teach_content": fzjjnr, "exhibit_class": zplb, "exhibition_class": zcfl,
                "exhibit_tech": zpjs, "method": jbczykxyyff, "effect": xxxg, "people_number": cgrs, "room": ztmc,
                "narrate_stru": ztxsjg, "weekend": xq, "record_time": jlsj, "record_people": jlr}
        data = {"label":"科技馆展品","venue":venueName,"item":item,"imgList":[],"videoList":[]}

        db.data.insert(data)
        # datas.append(data)
    # db.data.insert(datas)

